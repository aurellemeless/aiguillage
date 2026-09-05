import io
from typing import Optional

from docx import Document as DocxDocument
from fastapi import FastAPI, File, HTTPException, UploadFile
from pydantic import BaseModel

import generator

PROFILE_PATH = Path(__file__).resolve().parent.parent.parent / "profile" / "profile.json"

app = FastAPI(title="CV Generator", version="1.1.0")


class SkillLine(BaseModel):
    label: str
    values: list[str]


class Experience(BaseModel):
    company: str
    dates: str = ""
    role: str = ""
    bullets: list[str] = []
    tech: list[str] = []


class GenerateCVRequest(BaseModel):
    profile: dict
    output_name: str
    subdir: Optional[str] = None
    language: str = "fr"
    headline: str
    tagline: Optional[str] = None
    summary: Optional[str] = None
    skills: list[SkillLine]
    experience: list[Experience]
    personal_projects: Optional[list[Experience]] = None


class GenerateCoverLetterRequest(BaseModel):
    profile: dict
    output_name: str
    subdir: Optional[str] = None
    language: str = "fr"
    recipient: Optional[str] = None
    subject: Optional[str] = None
    body: list[str]


@app.get("/")
def health():
    return {"status": "ok"}


@app.post("/generate-cv")
def generate_cv(payload: GenerateCVRequest):
    content = payload.model_dump(exclude={"profile", "output_name", "subdir", "language"})
    doc = generator.build_cv(payload.profile, content, payload.language)
    path = generator.save_document(doc, payload.output_name, payload.subdir or "")
    return {"path": str(path)}


@app.post("/generate-cover-letter")
def generate_cover_letter(payload: GenerateCoverLetterRequest):
    content = payload.model_dump(exclude={"profile", "output_name", "subdir", "language"})
    doc = generator.build_cover_letter(payload.profile, content, payload.language)
    path = generator.save_document(doc, payload.output_name, payload.subdir or "")
    return {"path": str(path)}


@app.post("/extract-docx-text")
async def extract_docx_text(file: UploadFile = File(...)):
    raw = await file.read()
    try:
        doc = DocxDocument(io.BytesIO(raw))
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Fichier .docx illisible : {exc}")
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)
    return {"text": "\n".join(paragraphs)}
