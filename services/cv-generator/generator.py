"""Builds .docx files (CV and cover letter), reusing the visual system
documented in the `cv_generation_style` memory (originally implemented with
docx-js), ported here to python-docx.
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Twips

import style as st

DATA_DIR = Path(__file__).resolve().parent.parent.parent / "data" / "generated"

SECTIONS = {
    "fr": {
        "skills": "COMPÉTENCES",
        "experience": "EXPÉRIENCES PROFESSIONNELLES",
        "projects": "PROJETS PERSONNELS",
        "education": "FORMATIONS",
        "certifications": "CERTIFICATIONS",
        "languages": "LANGUES",
        "subject": "Objet : ",
        "tech": "Tech : ",
    },
    "en": {
        "skills": "SKILLS",
        "experience": "PROFESSIONAL EXPERIENCE",
        "projects": "PERSONAL PROJECTS",
        "education": "EDUCATION",
        "certifications": "CERTIFICATIONS",
        "languages": "LANGUAGES",
        "subject": "Subject: ",
        "tech": "Tech: ",
    },
}


def sections_for(language: str) -> dict:
    return SECTIONS.get(language, SECTIONS["fr"])


# ---------------------------------------------------------------------------
# Low-level OOXML helpers (docx-js exposed these as plain properties;
# python-docx requires manual XML manipulation for these cases).
# ---------------------------------------------------------------------------

def set_character_spacing(run, twips: int) -> None:
    rpr = run._element.get_or_add_rPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:val"), str(twips))
    rpr.append(spacing)


def set_paragraph_shading(paragraph, hex_color: str) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def set_paragraph_bottom_border(paragraph, hex_color: str, size: int = 8) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), hex_color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_cell_shading(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


# ---------------------------------------------------------------------------
# Document construction
# ---------------------------------------------------------------------------

def new_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Pt(st.MARGIN_TOP_BOTTOM_TWIPS / 20)
    section.bottom_margin = Pt(st.MARGIN_TOP_BOTTOM_TWIPS / 20)
    section.left_margin = Pt(st.MARGIN_LEFT_RIGHT_TWIPS / 20)
    section.right_margin = Pt(st.MARGIN_LEFT_RIGHT_TWIPS / 20)

    normal = doc.styles["Normal"]
    normal.font.name = st.FONT_NAME
    normal.font.size = Pt(st.SIZE_SUMMARY / 2)
    return doc


def add_run(paragraph, text, size, color, bold=False, italic=False):
    run = paragraph.add_run(text)
    run.font.name = st.FONT_NAME
    run.font.size = Pt(size / 2)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic
    return run


def add_header(doc: Document, identity: dict, headline: str, tagline: str | None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, headline, st.SIZE_TITLE, st.COLOR_TITLE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = add_run(p, identity["name"], st.SIZE_NAME, st.COLOR_NAME)
    set_character_spacing(run, st.CHAR_SPACING_NAME)

    if tagline:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, tagline, st.SIZE_TAGLINE, st.COLOR_TAGLINE, italic=True)

    contact = "   •   ".join(
        filter(
            None,
            [
                identity.get("email"),
                identity.get("phone"),
                identity.get("location"),
                identity.get("github"),
                identity.get("linkedin"),
            ],
        )
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, contact, st.SIZE_SUMMARY, st.COLOR_SUMMARY)


def add_summary(doc: Document, summary: str | None):
    if not summary:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, summary, st.SIZE_SUMMARY, st.COLOR_SUMMARY)


def add_section_header(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(180 / 20)
    p.paragraph_format.space_after = Pt(60 / 20)
    run = add_run(p, text, st.SIZE_SECTION_HEADER, st.COLOR_SECTION_HEADER, bold=True)
    set_character_spacing(run, st.CHAR_SPACING_SECTION_HEADER)
    set_paragraph_bottom_border(p, st.COLOR_SECTION_HEADER)


def add_skills_table(doc: Document, skills: list[dict]):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Twips(st.TABLE_COL_LABEL_TWIPS)
    table.columns[1].width = Twips(st.TABLE_COL_VALUE_TWIPS)

    for i, line in enumerate(skills):
        row = table.add_row()
        fill = st.COLOR_TABLE_ROW_ALT if i % 2 == 0 else st.COLOR_TABLE_ROW_BASE
        label_cell, value_cell = row.cells

        label_cell.width = Twips(st.TABLE_COL_LABEL_TWIPS)
        value_cell.width = Twips(st.TABLE_COL_VALUE_TWIPS)
        set_cell_shading(label_cell, fill)
        set_cell_shading(value_cell, fill)

        p = label_cell.paragraphs[0]
        add_run(p, line["label"], st.SIZE_TABLE_TEXT, st.COLOR_SECTION_HEADER, bold=True)

        p = value_cell.paragraphs[0]
        add_run(p, ", ".join(line["values"]), st.SIZE_TABLE_TEXT, st.COLOR_JOB_COMPANY)


def add_job_header(doc: Document, company: str, dates: str):
    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(Twips(st.JOB_HEADER_TAB_STOP_TWIPS), WD_TAB_ALIGNMENT.RIGHT)
    set_paragraph_shading(p, st.COLOR_JOB_SHADING)
    add_run(p, company, st.SIZE_JOB_COMPANY, st.COLOR_JOB_COMPANY, bold=True)
    p.add_run().add_tab()
    add_run(p, dates, st.SIZE_JOB_DATES, st.COLOR_JOB_DATES, italic=True)


def add_job_subtitle(doc: Document, role: str):
    p = doc.add_paragraph()
    add_run(p, role, st.SIZE_JOB_SUBTITLE, st.COLOR_JOB_SUBTITLE, italic=True)


def add_bullets(doc: Document, bullets: list[str]):
    for bullet in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(15 / 20)
        p.paragraph_format.space_after = Pt(15 / 20)
        add_run(p, bullet, st.SIZE_BULLET, st.COLOR_BULLET)


def add_tech_line(doc: Document, tech: list[str], tech_label: str):
    if not tech:
        return
    p = doc.add_paragraph()
    add_run(p, tech_label, st.SIZE_TECH, st.COLOR_TECH_LABEL, bold=True)
    add_run(p, ", ".join(tech), st.SIZE_TECH, st.COLOR_TECH_VALUE)


def add_experience_block(doc: Document, exp: dict, tech_label: str):
    add_job_header(doc, exp["company"], exp.get("dates", ""))
    if exp.get("role"):
        add_job_subtitle(doc, exp["role"])
    add_bullets(doc, exp.get("bullets", []))
    add_tech_line(doc, exp.get("tech", []), tech_label)


def add_simple_list_line(doc: Document, label: str, meta: str, date_str: str):
    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(Twips(st.JOB_HEADER_TAB_STOP_TWIPS), WD_TAB_ALIGNMENT.RIGHT)
    text = f"{label}  –  {meta}" if meta else label
    add_run(p, text, st.SIZE_JOB_DATES, st.COLOR_JOB_COMPANY)
    p.add_run().add_tab()
    add_run(p, date_str, st.SIZE_JOB_DATES, st.COLOR_JOB_DATES, italic=True)


def build_cv(profile: dict, content: dict, language: str = "fr") -> Document:
    doc = new_document()
    sections = sections_for(language)

    add_header(doc, profile["identity"], content["headline"], content.get("tagline"))
    add_summary(doc, content.get("summary"))

    add_section_header(doc, sections["skills"])
    add_skills_table(doc, content["skills"])

    add_section_header(doc, sections["experience"])
    for exp in content["experience"]:
        add_experience_block(doc, exp, sections["tech"])

    if content.get("personal_projects"):
        add_section_header(doc, sections["projects"])
        for project in content["personal_projects"]:
            add_experience_block(doc, project, sections["tech"])

    add_section_header(doc, sections["education"])
    for edu in profile["education"]:
        add_simple_list_line(doc, edu["degree"], edu["institution"], edu["year"])

    add_section_header(doc, sections["certifications"])
    for cert in profile["certifications"]:
        add_simple_list_line(doc, cert["name"], cert["organization"], cert["year"])

    add_section_header(doc, sections["languages"])
    p = doc.add_paragraph()
    languages_txt = "   •   ".join(f"{l['language']} – {l['level']}" for l in profile["languages"])
    if profile.get("driving_license"):
        languages_txt += f"   •   {profile['driving_license']}"
    add_run(p, languages_txt, st.SIZE_SUMMARY, st.COLOR_SUMMARY)

    return doc


def build_cover_letter(profile: dict, content: dict, language: str = "fr") -> Document:
    doc = new_document()
    identity = profile["identity"]
    sections = sections_for(language)

    p = doc.add_paragraph()
    add_run(p, identity["name"], st.SIZE_NAME, st.COLOR_NAME, bold=True)
    p = doc.add_paragraph()
    contact = "   •   ".join(
        filter(None, [identity.get("email"), identity.get("phone"), identity.get("location")])
    )
    add_run(p, contact, st.SIZE_TECH, st.COLOR_TECH_VALUE)

    doc.add_paragraph()

    if content.get("recipient"):
        p = doc.add_paragraph()
        add_run(p, content["recipient"], st.SIZE_SUMMARY, st.COLOR_SUMMARY)

    if content.get("subject"):
        p = doc.add_paragraph()
        add_run(p, f"{sections['subject']}{content['subject']}", st.SIZE_SUMMARY, st.COLOR_SUMMARY, bold=True)

    doc.add_paragraph()

    for paragraph_text in content["body"]:
        p = doc.add_paragraph()
        add_run(p, paragraph_text, st.SIZE_SUMMARY, st.COLOR_SUMMARY)
        doc.add_paragraph()

    return doc


def save_document(doc: Document, output_name: str, subdir: str = "") -> Path:
    target_dir = DATA_DIR / subdir if subdir else DATA_DIR
    target_dir.mkdir(parents=True, exist_ok=True)
    path = target_dir / output_name
    doc.save(path)
    return path
